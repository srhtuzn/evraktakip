import os
import sys, os, json, shutil, datetime, subprocess
import mysql.connector
from functools import partial
import docx
import random
import math
import openpyxl
import subprocess
import datetime as dt
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QDialog, QWidget, QTabWidget,
    QFormLayout, QHBoxLayout, QVBoxLayout, QSplitter,
    QLineEdit, QComboBox, QDateEdit, QPlainTextEdit,
    QTextEdit,
    QPushButton, QFileDialog, QMessageBox, QGroupBox,
    QTableWidget, QTableWidgetItem, QLabel, QCheckBox, QMenu, QInputDialog, QGraphicsView, QGraphicsScene, QGraphicsRectItem, QGraphicsTextItem, QCompleter
)
from PyQt6.QtCore import QDate, QTimer, Qt, QPointF
from PyQt6.QtGui import QIcon, QPalette, QColor, QFont, QPixmap, QPen, QPolygonF, QBrush, QPainterPath, QPainter
from PyQt6.QtPdfWidgets import QPdfView
from PyQt6.QtPdf import QPdfDocument



# Sabit şifre
DEFAULT_PASSWORD = "Ys123456"

def compute_username(name: str) -> str:
    parts = name.strip().split()
    if len(parts) >= 2:
        return (parts[0][0] + parts[-1]).lower()
    return parts[0].lower()

class SecureLoginDialog(QDialog):
    def __init__(self, conn_params):
        super().__init__()
        self.conn_params = conn_params
        self.authenticated = False
        self.current_user = None
        
        self.setWindowTitle("Evrak Takip - Giriş")
        self.setFixedSize(300, 200)
        
        layout = QVBoxLayout()
        
        # Kullanıcı Adı
        self.lbl_username = QLabel("Kullanıcı Adı:")
        self.txt_username = QLineEdit()
        self.txt_username.setPlaceholderText("kullanici_adi")
        layout.addWidget(self.lbl_username)
        layout.addWidget(self.txt_username)
        
        # Şifre
        self.lbl_password = QLabel("Şifre:")
        self.txt_password = QLineEdit()
        self.txt_password.setEchoMode(QLineEdit.EchoMode.Password)
        self.txt_password.setPlaceholderText("********")
        layout.addWidget(self.lbl_password)
        layout.addWidget(self.txt_password)
        
        # Giriş Butonu
        self.btn_login = QPushButton("Giriş Yap")
        self.btn_login.clicked.connect(self.authenticate)
        layout.addWidget(self.btn_login)
        
        # Hata Mesajı
        self.lbl_error = QLabel()
        self.lbl_error.setStyleSheet("color: red;")
        self.lbl_error.setVisible(False)
        layout.addWidget(self.lbl_error)
        
        self.setLayout(layout)
    
    def authenticate(self):
        username = self.txt_username.text().strip()
        password = self.txt_password.text().strip()
        
        if not username or not password:
            self.show_error("Kullanıcı adı ve şifre gereklidir!")
            return
            
        try:
            conn = mysql.connector.connect(**self.conn_params)
            cursor = conn.cursor(dictionary=True)
            
            cursor.execute("""
                SELECT id, username, name 
                FROM users 
                WHERE username = %s AND password = %s
            """, (username, password))
            
            user = cursor.fetchone()
            conn.close()
            
            if user:
                self.authenticated = True
                self.current_user = user
                self.accept()  # Başarılı giriş
            else:
                self.show_error("Geçersiz kullanıcı adı veya şifre!")
                
        except Exception as e:
            self.show_error(f"Veritabanı hatası: {str(e)}")
    
    def show_error(self, message):
        self.lbl_error.setText(message)
        self.lbl_error.setVisible(True)
        QTimer.singleShot(3000, lambda: self.lbl_error.setVisible(False))


    def check_credentials(self):
        u = self.user_le.text().strip()
        p = self.pw_le.text().strip()
        if not u or not p:
            QMessageBox.warning(self, "Hata", "Kullanıcı adı ve şifre boş olamaz.")
            return
        try:
            conn = mysql.connector.connect(**self.conn_params)
            cur  = conn.cursor()
            cur.execute("SELECT password FROM users WHERE username=%s", (u,))
            row = cur.fetchone()
            cur.close()
            conn.close()
            if not row or row[0] != p:
                QMessageBox.critical(self, "Hata", "Geçersiz kullanıcı adı veya şifre.")
            else:
                self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Veritabanı hatası:\n{e}")

class MainWindow(QMainWindow):
    def __init__(self, cfg, conn_params, current_user):
        super().__init__()
        self.current_user = current_user  # Artık tüm kullanıcı bilgileri mevcut
        print(f"Oturum açan kullanıcı: {self.current_user['name']}")
        self.edit_mode = False
        self.edit_tracking_no = None
        self.current_doc_id = None
        self.cfg = cfg
        self.conn_params = conn_params

        self.responsibles = []
        self.companies     = []
        self.projects      = []
        self.subjects      = []
        self.all_docs      = []

        self.init_ui()
        self.load_lookups()
        self.refresh_lookup_widgets()

        self.page_size    = 50
        self.current_page = 0
        self.total_records= 0
        self.scheduler    = BackgroundScheduler()
        self.scheduler.start()

        # >>> Ortak araç çubuğuna yenile butonu
        refresh_btn = QPushButton("🔄 Yenile")
        refresh_btn.clicked.connect(self.refresh_all)
        self.tool_bar = self.addToolBar("Araçlar")
        self.tool_bar.addWidget(refresh_btn)
        self.setStyleSheet("""
        QComboBox {
            background-color: white;
            color: black;
            selection-background-color: #8e44ad;
            selection-color: white;
        }
        QComboBox QAbstractItemView {
            background-color: white;
            color: black;
            selection-background-color: #8e44ad;
            selection-color: white;
        }
    """)


    def refresh_all(self):
        self.load_lookups()
        self.refresh_lookup_widgets()

    def db_connect(self):
        return mysql.connector.connect(**self.conn_params)

    def load_lookups(self):
        conn = self.db_connect()
        cur  = conn.cursor()
        # Sorumlular
        cur.execute("SELECT id, name FROM users")
        self.responsibles = cur.fetchall()
        # Şirketler
        cur.execute("SELECT id, name FROM companies")
        self.companies = cur.fetchall()
        # Projeler (id, name, company_id)
        cur.execute("SELECT id, name, company_id FROM projects")
        self.projects = cur.fetchall()
        # Konular (artık project_id de var)
        cur.execute("SELECT id, name, project_id FROM subjects ORDER BY name")
        self.subjects = cur.fetchall()  # liste öğesi: (sid, sname, proj_id)
        # Tüm evraklar (İlgili Belge için archive_name)
        cur.execute("SELECT id, archive_name FROM documents ORDER BY created_at DESC")
        self.all_docs = cur.fetchall()
        cur.close()
        conn.close()



    def refresh_lookup_widgets(self):
        # --- Evrak Kaydı sekmesi ---
        # Sorumlu
        self.resp_cb.clear()
        for uid, uname in self.responsibles:
            self.resp_cb.addItem(uname, uid)
        # Şirket
        self.company_cb.clear()
        self.company_cb.addItem("", None)
        for cid, cname in self.companies:
            self.company_cb.addItem(cname, cid)
        # Proje
        self.project_cb.clear()
        self.project_cb.addItem("", None)
        for pid, pname, comp_id in self.projects:
            self.project_cb.addItem(pname, pid)
        # Konu — başta tüm konuları listele, sonra project değişince filtreleceğiz
        self.subject_cb.clear()
        self.subject_cb.addItem("— Seçiniz —", None)
        for sid, sname, proj_id in self.subjects:
            self.subject_cb.addItem(sname, sid)
        # İlgili Belge (arama destekli)
        self.link_parent_cb.clear()
        self.link_parent_cb.setEditable(True)
        self.link_parent_cb.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.link_parent_cb.addItem("— Yok —", None)
        for doc_id, aname in self.all_docs:
            self.link_parent_cb.addItem(aname, doc_id)
        completer = QCompleter([aname for _, aname in self.all_docs], self.link_parent_cb)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.link_parent_cb.setCompleter(completer)

        # --- Evrak Ara sekmesi ---
        self.search_doc_type_cb.clear()
        self.search_doc_type_cb.addItem("", None)
        self.search_doc_type_cb.addItem("Resmi Evrak",   "official")
        self.search_doc_type_cb.addItem("Taşınmaz Evrak", "immovable")
        self.search_doc_type_cb.addItem("Proje",          "project")
        self.search_company_cb.clear()
        self.search_company_cb.addItem("", None)
        for cid, cname in self.companies:
            self.search_company_cb.addItem(cname, cid)
        self.search_project_cb.clear()
        self.search_project_cb.addItem("", None)
        for pid, pname, comp_id in self.projects:
            self.search_project_cb.addItem(pname, pid)
        self.search_res_cb.clear()
        self.search_res_cb.addItem("", None)
        for uid, uname in self.responsibles:
            self.search_res_cb.addItem(uname, uid)

        # --- Yönetim sekmesi ---
        self.del_res_cb.clear()
        for uid, uname in self.responsibles:
            self.del_res_cb.addItem(uname, uid)
        self.del_company_cb.clear()
        for cid, cname in self.companies:
            self.del_company_cb.addItem(cname, cid)
        self.add_proj_company_cb.clear()
        for cid, cname in self.companies:
            self.add_proj_company_cb.addItem(cname, cid)
        self.del_project_cb.clear()
        for pid, pname, comp_id in self.projects:
            self.del_project_cb.addItem(pname, pid)
        # Konu silme
        self.del_subj_cb.clear()
        for sid, sname, proj_id in self.subjects:
            self.del_subj_cb.addItem(sname, sid)

        # --- İş Akışı sekmesi (Eksik olan burasıydı) ---
        self.combo_company_wf.clear()
        self.combo_company_wf.addItem("— Seçiniz —", None)
        for cid, cname in self.companies:
            self.combo_company_wf.addItem(cname, cid)

        self.combo_project_wf.clear()
        self.combo_project_wf.addItem("— Seçiniz —", None)

        self.combo_topic_wf.clear()
        self.combo_topic_wf.addItem("— Seçiniz —", None)



    def clear_form(self):
        self.fetch_tracking_le.clear()
        self.file_path_le.clear()
        self.archive_name_le.clear()
        self.subject_cb.setCurrentIndex(0)
        self.doc_type_cb.setCurrentIndex(0)
        self.doc_date_de.setDate(QDate.currentDate())
        self.resp_cb.setCurrentIndex(0)
        self.sender_le.clear()
        self.recipient_le.clear()
        self.company_cb.setCurrentIndex(0)
        self.project_cb.setCurrentIndex(0)
        self.start_date_de.setDate(QDate.currentDate())
        self.end_date_de.setDate(QDate.currentDate())
        self.comment_te.clear()
        self.link_parent_cb.setCurrentIndex(0)
        self.target_folder_le.clear()

        # edit modunu kapat ve buton metnini eski haline getir
        self.edit_mode = False
        self.edit_tracking_no = None
        self.save_btn.setText("Kaydet")

    def on_project_changed_update_subjects(self):
        pid = self.project_cb.currentData()
        self.subject_cb.blockSignals(True)
        self.subject_cb.clear()

        if pid is None:
            self.subject_cb.addItem("Önce proje seçin", None)
            self.subject_cb.setEnabled(False)
        else:
            self.subject_cb.setEnabled(True)
            self.subject_cb.addItem("Seçiniz", None)
            for sid, sname, proj_id in self.subjects:
                if proj_id == pid:
                    self.subject_cb.addItem(sname, sid)

        self.subject_cb.blockSignals(False)

    def on_company_changed_update_projects(self):
        cid = self.company_cb.currentData()
        self.project_cb.blockSignals(True)
        self.project_cb.clear()
        self.subject_cb.clear()
        self.subject_cb.setEnabled(False)

        if cid is None:
            self.project_cb.addItem("Önce şirket seçin", None)
            self.project_cb.setEnabled(False)
        else:
            self.project_cb.setEnabled(True)
            self.project_cb.addItem("Seçiniz", None)
            for pid, pname, comp_id in self.projects:
                if comp_id == cid:
                    self.project_cb.addItem(pname, pid)

        self.project_cb.blockSignals(False)




    def generate_tracking_no(self):
        today = dt.datetime.now().strftime("%Y-%m-%d")

        conn = self.db_connect()
        cur = conn.cursor()

        for _ in range(5):  # Maksimum 5 kez dene
            suffix = random.randint(1, 999)
            tn = f"{today}-{suffix:03}"

            cur.execute("SELECT 1 FROM documents WHERE tracking_number = %s", (tn,))
            if not cur.fetchone():
                cur.close()
                conn.close()
                return tn  # benzersizse döndür

        cur.close()
        conn.close()
        raise ValueError("Takip numarası üretilemedi. Lütfen tekrar deneyin.")

    def init_ui(self):
        self.setWindowTitle("Evrak Takip")
        tabs = QTabWidget()
        tabs.addTab(self.build_document_tab(), "Evrak Kaydı")
        tabs.addTab(self.build_search_tab(),   "Evrak Bul")
        tabs.addTab(self.build_workflow_tab(), "İş Akışı")
        tabs.addTab(self.build_notifications_tab(), "Bildirimler")
        tabs.addTab(self.build_manage_tab(),   "Yönetim")
        tabs.addTab(self.build_backup_tab(),   "Yedekleme")
        self.setCentralWidget(tabs)

        # — İşte buraya ekle —
        self.time_label = QLabel()
        self.statusBar().addPermanentWidget(self.time_label)

        timer = QTimer(self)
        timer.timeout.connect(self.update_clock)
        timer.start(1000)
        self.update_clock()
    def update_clock(self):
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.time_label.setText(now)


    def on_wf_company_changed(self, _):
        sel_cid = self.combo_company_wf.currentData()
        self.combo_project_wf.clear()
        self.combo_project_wf.addItem("— Tümü —", None)
        for pid, pname, comp_id in self.projects:
            if sel_cid is None or comp_id == sel_cid:
                self.combo_project_wf.addItem(pname, pid)
        # Konu (topic) combobox’unu temizle
        self.combo_topic_wf.clear()
        self.combo_topic_wf.addItem("— Seçiniz —", None)


    def on_wf_project_changed(self, _):
        sel_pid = self.combo_project_wf.currentData()
        self.combo_topic_wf.clear()
        self.combo_topic_wf.addItem("— Seçiniz —", None)
        if sel_pid is None:
            return
        conn = self.db_connect(); cur = conn.cursor()
        # Konu combobox’ını seçilen projeye bağlı subjects tablosundan al
        cur.execute("""
            SELECT id, name
              FROM subjects
             WHERE project_id = %s
             ORDER BY name
        """, (sel_pid,))
        for sid, sname in cur.fetchall():
            self.combo_topic_wf.addItem(sname, sid)
        cur.close(); conn.close()

    def build_workflow_tab(self):
        container = QWidget()
        # 1) Önce layout’u tanımla:
        layout = QVBoxLayout(container)
        # 2) Sonra form layout’u oluştur ve ekle
        form = QFormLayout()
        layout.addLayout(form)
        # 1) Şirket / Proje / Konu seçim combobox’ları
        self.combo_company_wf = QComboBox(); form.addRow("Şirket", self.combo_company_wf)
        self.combo_project_wf = QComboBox(); form.addRow("Proje",  self.combo_project_wf)
        self.combo_topic_wf   = QComboBox(); form.addRow("Konu",   self.combo_topic_wf)

        # 2) Seçimler değişince akışı yeniden yükle
        self.combo_company_wf.currentIndexChanged.connect(self.on_wf_company_changed)
        self.combo_project_wf.currentIndexChanged.connect(self.on_wf_project_changed)
        self.combo_topic_wf.currentIndexChanged.connect(self.load_workflow)

        # 3) Grafik sahnesi & view
        self.wf_scene = QGraphicsScene(self)
        self.wf_view  = QGraphicsView(self.wf_scene, container)
        self.wf_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        # Çift tıklayınca detay penceresi açsın
        self.wf_view.mouseDoubleClickEvent = self.show_details
        
        # 4) Layout
        layout.addWidget(self.wf_view)
        return container

    def load_workflow(self):
        # konu seçilmemişse temizle
        subject_id = self.combo_topic_wf.currentData()
        self.wf_scene.clear()
        if subject_id is None:
            return

        # 1) Başlangıç evrakları al
        conn, cur = self.db_connect(), None
        try:
            cur = conn.cursor()
            cur.execute("SELECT id FROM documents WHERE subject_id=%s", (subject_id,))
            initial_docs = [r[0] for r in cur.fetchall()]

            # 2) BFS ile tüm bağlantıları topla
            to_visit, visited, edges, nodes = initial_docs[:], set(), [], set()
            while to_visit:
                did = to_visit.pop()
                if did in visited: continue
                visited.add(did); nodes.add(did)
                for dir_, q in [("parent_id", "child_id"), ("child_id","parent_id")]:
                    cur.execute(f"SELECT {q} FROM document_links WHERE {dir_}=%s", (did,))
                    for (other,) in cur.fetchall():
                        edges.append((did if q=="child_id" else other,
                                      other if q=="child_id" else did))
                        to_visit.append(other)

            # 3) Node bilgilerini çek
            self.id_to_info = {}
            if nodes:
                cur.execute(
                    "SELECT id, tracking_number, archive_name, doc_type, end_date "
                    "FROM documents WHERE id IN (%s)" % ",".join(map(str, nodes))
                )
                for _id, tn, an, dt, ed in cur.fetchall():
                    label = f"{tn}\n{an}\n{dt}\n{ed.strftime('%d.%m.%Y') if ed else ''}"
                    details = f"Takip No: {tn}\nAdı: {an}\nTip: {dt}\nBitiş: {ed or '-'}"
                    self.id_to_info[_id] = {"label": label, "details": details}

            # 4) Seviyelendir ve çiz
            level_map = {}
            def assign(nid, lvl):
                level_map[nid] = max(level_map.get(nid, lvl), lvl)
                for (p,c) in edges:
                    if p==nid: assign(c, lvl+1)
            for root in initial_docs:
                assign(root, 0)

            lvl_nodes = {}
            for nid, lvl in level_map.items():
                lvl_nodes.setdefault(lvl, []).append(nid)

            rect_w, base_h, x_sp, y_sp = 160, 30, 200, 120
            node_pos = {}
            for lvl, ids in lvl_nodes.items():
                for i, nid in enumerate(ids):
                    x, y = i*x_sp, lvl*y_sp
                    text = self.id_to_info[nid]['label']
                    lines = text.count("\n")+1
                    h = lines*18 + 10
                    rect = QGraphicsRectItem(x, y, rect_w, h)
                    rect.setBrush(QBrush(QColor("#d0eaff")))
                    rect.setData(0, nid)
                    self.wf_scene.addItem(rect)
                    ti = QGraphicsTextItem(text, rect)
                    ti.setTextWidth(rect_w-4)
                    ti.setPos(x+2, y+2)
                    node_pos[nid] = (x,y,h)

            pen = QPen(Qt.GlobalColor.black)
            for src,dst in edges:
                if src in node_pos and dst in node_pos:
                    x1,y1,h1 = node_pos[src]
                    x2,y2,h2 = node_pos[dst]
                    start = QPointF(x1+rect_w/2, y1+h1)
                    end   = QPointF(x2+rect_w/2, y2)
                    path  = QPainterPath(start)
                    path.lineTo(end)
                    self.wf_scene.addPath(path, pen)
                    # Ok başı
                    sz = 6
                    ang = math.atan2(end.y()-start.y(), end.x()-start.x())
                    p1  = end + QPointF(-sz*math.cos(ang-math.pi/6), -sz*math.sin(ang-math.pi/6))
                    p2  = end + QPointF(-sz*math.cos(ang+math.pi/6), -sz*math.sin(ang+math.pi/6))
                    self.wf_scene.addPolygon(QPolygonF([end,p1,p2]), pen, QBrush(Qt.GlobalColor.black))
        finally:
            if cur: cur.close()
            conn.close()

    def show_details(self, event):
        scene_pt = self.wf_view.mapToScene(event.pos())
        item     = self.wf_scene.itemAt(scene_pt, self.wf_view.transform())
        # eğer TextItem’ın içindeyse, parent rect’e çık
        if isinstance(item, QGraphicsTextItem):
            item = item.parentItem()
        if isinstance(item, QGraphicsRectItem):
            nid = item.data(0)
            if nid in self.id_to_info:
                QMessageBox.information(self, "Evrak Detayı",
                                        self.id_to_info[nid]['details'])







    def build_document_tab(self):
        form = QFormLayout()

        # 4) Evrak Getir
        self.fetch_tracking_le = QLineEdit()
        fetch_btn = QPushButton("Evrak Getir")
        fetch_btn.clicked.connect(self.fetch_document_by_tracking_no)
        hb_fetch = QHBoxLayout()
        hb_fetch.addWidget(self.fetch_tracking_le)
        hb_fetch.addWidget(fetch_btn)
        form.addRow("Takip No ile Getir", hb_fetch)

        # 5) Evrak Dosyası *
        self.file_path_le = QLineEdit()
        self.file_path_le.setReadOnly(True)
        sel_file_btn = QPushButton("Dosya Seç")
        sel_file_btn.clicked.connect(self.choose_file)
        hb = QHBoxLayout()
        hb.addWidget(self.file_path_le)
        hb.addWidget(sel_file_btn)
        box = QWidget(); box.setLayout(hb)
        form.addRow("Evrak Dosyası *", box)

        # 1) Şirket *
        self.company_cb = QComboBox()
        self.company_cb.addItem("Seçiniz", None)
        for cid, cname in self.companies:
            self.company_cb.addItem(cname, cid)
        self.company_cb.currentIndexChanged.connect(self.on_company_changed_update_projects)
        form.addRow("Şirket *", self.company_cb)

        # 2) Proje *
        self.project_cb = QComboBox()
        self.project_cb.setEnabled(False)
        self.project_cb.addItem("Önce şirket seçin", None)
        self.project_cb.currentIndexChanged.connect(self.on_project_changed_update_subjects)
        form.addRow("Proje *", self.project_cb)

        # 3) Konu *
        self.subject_cb = QComboBox()
        self.subject_cb.setEnabled(False)
        self.subject_cb.addItem("Önce proje seçin", None)
        form.addRow("Konu *", self.subject_cb)


        # 6) Arşiv Adlandırması *
        self.archive_name_le = QLineEdit()
        form.addRow("Arşiv Adlandırması *", self.archive_name_le)

        # 7) Evrak Tipi *
        self.doc_type_cb = QComboBox()
        self.doc_type_cb.addItem("Resmi Evrak",   "official")
        self.doc_type_cb.addItem("Taşınmaz Evrak", "immovable")
        self.doc_type_cb.addItem("Proje",          "project")
        form.addRow("Evrak Tipi *", self.doc_type_cb)

        # 8) Evrak Tarihi *
        self.doc_date_de = QDateEdit(QDate.currentDate())
        self.doc_date_de.setCalendarPopup(True)
        form.addRow("Evrak Tarihi *", self.doc_date_de)

        # 9) Sorumlu *
        self.resp_cb = QComboBox()
        form.addRow("Sorumlu *", self.resp_cb)

        # 10) Gönderen *
        self.sender_le = QLineEdit()
        form.addRow("Gönderen *", self.sender_le)

        # 11) Alıcı *
        self.recipient_le = QLineEdit()
        form.addRow("Alıcı *", self.recipient_le)

        # 12) Başlangıç / Bitiş
        self.start_date_de = QDateEdit(QDate.currentDate()); self.start_date_de.setCalendarPopup(True)
        self.end_date_de   = QDateEdit(QDate.currentDate()); self.end_date_de.setCalendarPopup(True)
        form.addRow("Başlangıç Tarihi", self.start_date_de)
        form.addRow("Bitiş Tarihi",      self.end_date_de)

        # 13) Yorum
        self.comment_te = QPlainTextEdit()
        form.addRow("Yorum", self.comment_te)

        # 14) İlgili Belge (opsiyonel)
        self.link_parent_cb = QComboBox()
        self.link_parent_cb.addItem("— Yok —", None)
        form.addRow("İlgili Belge (opsiyonel)", self.link_parent_cb)

        # 15) Kaydet / Güncelle
        self.save_btn = QPushButton("Kaydet")
        self.save_btn.clicked.connect(self.on_save)
        form.addRow(self.save_btn)

        # 16) Kayıt Klasörü *
        self.target_folder_le = QLineEdit(); self.target_folder_le.setReadOnly(True)
        sel_folder_btn = QPushButton("Klasör Seç")
        sel_folder_btn.clicked.connect(self.choose_target_folder)
        hb2 = QHBoxLayout()
        hb2.addWidget(self.target_folder_le)
        hb2.addWidget(sel_folder_btn)
        box2 = QWidget(); box2.setLayout(hb2)
        form.addRow("Kayıt Klasörü *", box2)

        w = QWidget()
        w.setLayout(form)
        return w
    
    def fetch_document_by_tracking_no(self):
        tn = self.fetch_tracking_le.text().strip()
        if not tn:
            QMessageBox.warning(self, "Uyarı", "Lütfen takip numarasını girin.")
            return

        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("""
            SELECT archive_name, subject, subject_id, doc_type, doc_date, responsible, sender, recipient,
                company, project, start_date, end_date, comment, file_path
            FROM documents
            WHERE tracking_number = %s
        """, (tn,))
        row = cur.fetchone()
        cur.close()
        conn.close()

        if not row:
            QMessageBox.warning(self, "Bulunamadı", f"{tn} numaralı evrak bulunamadı.")
            return

        self.archive_name_le.setText(row[0])
        self.subject_cb.setCurrentIndex(self.subject_cb.findData(row[2]))  # subject_id
        self.doc_type_cb.setCurrentIndex(self.doc_type_cb.findData(row[3]))
        self.doc_date_de.setDate(QDate.fromString(row[4].strftime("%Y-%m-%d"), "yyyy-MM-dd"))
        self.resp_cb.setCurrentIndex(self.resp_cb.findData(row[5]))
        self.sender_le.setText(row[6])
        self.recipient_le.setText(row[7])
        self.company_cb.setCurrentIndex(self.company_cb.findData(row[8]))
        self.project_cb.setCurrentIndex(self.project_cb.findData(row[9]))

        if row[10]:
            self.start_date_de.setDate(QDate.fromString(row[10].strftime("%Y-%m-%d"), "yyyy-MM-dd"))
        if row[11]:
            self.end_date_de.setDate(QDate.fromString(row[11].strftime("%Y-%m-%d"), "yyyy-MM-dd"))

        self.comment_te.setPlainText(row[12] or "")
        self.file_path_le.setText(row[13])
        self.edit_mode = True
        self.edit_tracking_no = tn
        self.save_btn.setText("Güncelle")


    
    def build_backup_tab(self):
        form = QFormLayout()

        # Yedek klasörü
        self.backup_folder_le = QLineEdit(); self.backup_folder_le.setReadOnly(True)
        sel_backup_btn = QPushButton("Klasör Seç"); sel_backup_btn.clicked.connect(self.choose_backup_folder)
        hb = QHBoxLayout(); hb.addWidget(self.backup_folder_le); hb.addWidget(sel_backup_btn)
        box = QWidget(); box.setLayout(hb)
        form.addRow("Yedek Klasörü *", box)

        # Tam yedek seçeneği
        self.full_backup_cb = QCheckBox("Tüm veritabanları ve izinlerle yedekle")
        form.addRow(self.full_backup_cb)

        # Frekans seçimi
        self.backup_freq_cb = QComboBox()
        self.backup_freq_cb.addItems(["Günlük","Haftalık","Aylık"])
        form.addRow("Frekans", self.backup_freq_cb)

        # Zaman seçimi (saat:dakika)
        self.backup_time_le = QLineEdit("02:00")  
        form.addRow("Saat (HH:MM)", self.backup_time_le)

        # Planla butonu
        plan_btn = QPushButton("Planla")
        plan_btn.clicked.connect(self.schedule_backup)
        form.addRow(plan_btn)

        # Anında yedekleme butonu
        now_btn = QPushButton("Hemen Yedekle")
        now_btn.clicked.connect(self.do_backup)
        form.addRow(now_btn)

        w = QWidget(); w.setLayout(form)
        return w
    def schedule_backup(self):
        folder = self.backup_folder_le.text().strip()
        if not os.path.isdir(folder):
            QMessageBox.warning(self, "Hata", "Lütfen geçerli bir yedek klasörü seçin.")
            return

        freq = self.backup_freq_cb.currentText()
        time_str = self.backup_time_le.text().strip()
        try:
            hour, minute = map(int, time_str.split(":"))
        except:
            QMessageBox.warning(self, "Hata", "Saat formatı HH:MM olmalı.")
            return

        # Önce varsa eski job'ı temizle
        if hasattr(self, "_backup_job_id"):
            try: self.scheduler.remove_job(self._backup_job_id)
            except: pass

        # CronTrigger ayarla
        if freq == "Günlük":
            trigger = CronTrigger(hour=hour, minute=minute)
        elif freq == "Haftalık":
            trigger = CronTrigger(day_of_week="mon", hour=hour, minute=minute)
        else:  # Aylık
            trigger = CronTrigger(day=1, hour=hour, minute=minute)

        # Job ekle
        job = self.scheduler.add_job(
            func=self.do_backup,
            trigger=trigger,
            args=[],
            id=f"db_backup_{freq.lower()}",
            replace_existing=True
        )
        self._backup_job_id = job.id

        QMessageBox.information(
            self, "Planlandı",
            f"{freq} yedekleme saat {time_str} için planlandı."
        )


    def choose_backup_folder(self):
        # Başlangıç olarak kullanıcı dizini
        start = os.path.expanduser("~")
        path = QFileDialog.getExistingDirectory(
            self,
            "Yedek Klasörü Seçin",
            start,
            QFileDialog.Option.ShowDirsOnly | QFileDialog.Option.DontUseNativeDialog
        )
        if path:
            self.backup_folder_le.setText(path)
    def do_backup(self):
            # 1) Hedef klasörü al ve doğrula
            folder = self.backup_folder_le.text().strip()
            if not os.path.isdir(folder):
                QMessageBox.warning(self, "Hata", "Lütfen geçerli bir yedek klasörü seçin.")
                return

            # 2) Zaman damgasıyla yedek dosyası adı oluştur
            ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"evrak_takip_backup_{ts}.sql"
            dest = os.path.join(folder, filename)

            # 3) mysqldump komutunu hazırla
            mc = self.cfg['mysql']
            cmd = [
                "mysqldump",
                "--routines",   # Stored procedures & functions
                "--events",     # Scheduled events
                "-h", mc['host'],
                "-P", str(mc['port']),
                "-u", mc['user'],
                f"-p{mc['password']}",
                "evrak_takip"   # Sadece bu veritabanı adı
            ]

            # 4) Komutu çalıştır ve sonucu dosyaya yaz
            try:
                with open(dest, "w", encoding="utf8") as f:
                    subprocess.run(cmd, stdout=f, check=True)
                QMessageBox.information(
                    self, "Tamam",
                    f"`evrak_takip` veritabanı yedeği başarıyla alındı:\n{dest}"
                )
            except subprocess.CalledProcessError as e:
                QMessageBox.critical(
                    self, "Hata",
                    f"Yedekleme sırasında hata oluştu:\n{e}"
                )

    def build_search_tab(self):
        # Splitter: sol filtre+tablo, sağ preview paneli
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Sol taraf: filtre form + sonuç tablosu ---
        left_w = QWidget()
        left_layout = QVBoxLayout(left_w)

        form = QFormLayout()
        self.search_archive_le  = QLineEdit()
        form.addRow("Arşiv Adlandırması", self.search_archive_le)
        self.search_tracking_le = QLineEdit()
        form.addRow("Takip No",           self.search_tracking_le)

        self.search_doc_type_cb = QComboBox()
        form.addRow("Evrak Tipi", self.search_doc_type_cb)

        self.search_company_cb = QComboBox()
        form.addRow("Şirket", self.search_company_cb)

        self.search_project_cb = QComboBox()
        form.addRow("Proje", self.search_project_cb)

        self.search_res_cb = QComboBox()
        form.addRow("Sorumlu", self.search_res_cb)

        self.search_date_from = QDateEdit(QDate.currentDate())
        self.search_date_from.setCalendarPopup(True)
        self.search_date_to = QDateEdit(QDate.currentDate())
        self.search_date_to.setCalendarPopup(True)
        form.addRow("Tarih Başlangıç", self.search_date_from)
        form.addRow("Tarih Bitiş",      self.search_date_to)

        btn = QPushButton("Ara")
        btn.clicked.connect(self.search_records)
        form.addRow(btn)

        left_layout.addLayout(form)

        # --- Tablo ---
        self.results_table = QTableWidget()
        headers = ["Takip No","Arşiv Ad.","Tip","Şirket","Proje","Tarih","Sorumlu","Gönderen","Alıcı","Yorum"]
        self.results_table.setColumnCount(len(headers))
        self.results_table.setHorizontalHeaderLabels(headers)

        # Sağ-tık menüsü için hem tabloya hem viewport’a politika ve sinyal bağla
        self.results_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.results_table.customContextMenuRequested.connect(self.open_context_menu)
        self.results_table.viewport().setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.results_table.viewport().customContextMenuRequested.connect(self.open_context_menu)

        left_layout.addWidget(self.results_table)

        # Sayfalama kontrolleri
        nav_hb = QHBoxLayout()
        self.prev_btn = QPushButton("◀ Önceki")
        self.next_btn = QPushButton("Sonraki ▶")
        self.page_lbl = QLabel("Sayfa: 0 / 0")
        self.prev_btn.clicked.connect(self.on_prev_page)
        self.next_btn.clicked.connect(self.on_next_page)
        nav_hb.addWidget(self.prev_btn)
        nav_hb.addWidget(self.page_lbl)
        nav_hb.addWidget(self.next_btn)
        left_layout.addLayout(nav_hb)

        splitter.addWidget(left_w)

        # --- Sağ taraf: önizleme paneli ---
        self.preview_stack = QWidget()
        preview_layout = QVBoxLayout(self.preview_stack)

        # PDF önizleme
        self.pdf_doc  = QPdfDocument(self)
        self.pdf_view = QPdfView(self.preview_stack)
        self.pdf_view.setDocument(self.pdf_doc)
        preview_layout.addWidget(self.pdf_view)
        self.pdf_view.hide()

        # DOCX metin önizleme
        self.text_view = QTextEdit(self.preview_stack)
        self.text_view.setReadOnly(True)
        preview_layout.addWidget(self.text_view)
        self.text_view.hide()

        # Excel tablo önizleme
        self.xlsx_table = QTableWidget(self.preview_stack)
        preview_layout.addWidget(self.xlsx_table)
        self.xlsx_table.hide()

        # *** Görsel önizleme için QLabel ***
        self.image_view = QLabel(self.preview_stack)
        self.image_view.setAlignment(Qt.AlignmentFlag.AlignCenter)
        preview_layout.addWidget(self.image_view)
        self.image_view.hide()

        splitter.addWidget(self.preview_stack)

        # Boyutlama ve sabitleme
        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 4)
        splitter.setCollapsible(0, False)
        splitter.setCollapsible(1, False)
        self.preview_stack.setMinimumWidth(300)

        # Container
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.addWidget(splitter)
        return container


    def build_manage_tab(self):
        self.load_lookups()  # Şirket/proje/konu listelerini güncelle

        layout = QVBoxLayout()

        # Sorumlu Yönetimi
        grp1 = QGroupBox("Sorumlu Kişi Yönetimi")
        h1 = QHBoxLayout()
        self.add_res_le = QLineEdit()
        h1.addWidget(self.add_res_le)
        btn1 = QPushButton("Ekle")
        btn1.clicked.connect(self.add_responsible)
        h1.addWidget(btn1)
        self.del_res_cb = QComboBox()
        h1.addWidget(self.del_res_cb)
        btn1d = QPushButton("Sil")
        btn1d.clicked.connect(self.delete_responsible)
        h1.addWidget(btn1d)
        grp1.setLayout(h1)
        layout.addWidget(grp1)

        # Şirket Yönetimi
        grp2 = QGroupBox("Şirket Yönetimi")
        h2 = QHBoxLayout()
        self.add_company_le = QLineEdit()
        h2.addWidget(self.add_company_le)
        btn2 = QPushButton("Ekle")
        btn2.clicked.connect(self.add_company)
        h2.addWidget(btn2)
        self.del_company_cb = QComboBox()
        h2.addWidget(self.del_company_cb)
        btn2d = QPushButton("Sil")
        btn2d.clicked.connect(self.delete_company)
        h2.addWidget(btn2d)
        grp2.setLayout(h2)
        layout.addWidget(grp2)

        # Proje Yönetimi
        grp3 = QGroupBox("Proje Yönetimi")
        h3 = QHBoxLayout()
        self.add_proj_le = QLineEdit()
        h3.addWidget(self.add_proj_le)
        self.add_proj_company_cb = QComboBox()
        h3.addWidget(self.add_proj_company_cb)
        btn3 = QPushButton("Ekle")
        btn3.clicked.connect(self.add_project)
        h3.addWidget(btn3)
        self.del_project_cb = QComboBox()
        h3.addWidget(self.del_project_cb)
        btn3d = QPushButton("Sil")
        btn3d.clicked.connect(self.delete_project)
        h3.addWidget(btn3d)
        grp3.setLayout(h3)
        layout.addWidget(grp3)

        # Konu Yönetimi
        grp4 = QGroupBox("Konu Yönetimi")
        h4 = QHBoxLayout()
        self.add_subj_le = QLineEdit()
        self.add_subj_le.setPlaceholderText("Yeni konu adı")
        h4.addWidget(self.add_subj_le)

        self.add_subj_company_cb = QComboBox()
        self.add_subj_company_cb.addItem("— Şirket Seçiniz —", None)
        for cid, cname in self.companies:
            self.add_subj_company_cb.addItem(cname, cid)
        self.add_subj_company_cb.currentIndexChanged.connect(self.on_add_subj_company_changed)
        h4.addWidget(self.add_subj_company_cb)

        self.add_subj_proj_cb = QComboBox()
        self.add_subj_proj_cb.addItem("— Proje Seçiniz —", None)
        h4.addWidget(self.add_subj_proj_cb)

        btn_add_subj = QPushButton("Ekle")
        btn_add_subj.clicked.connect(self.add_subject)
        h4.addWidget(btn_add_subj)

        self.del_subj_cb = QComboBox()
        h4.addWidget(self.del_subj_cb)
        btn_del_subj = QPushButton("Sil")
        btn_del_subj.clicked.connect(self.delete_subject)
        h4.addWidget(btn_del_subj)

        grp4.setLayout(h4)
        layout.addWidget(grp4)

        w = QWidget()
        w.setLayout(layout)
        return w


    def on_add_subj_company_changed(self, index):
        cid = self.add_subj_company_cb.currentData()
        self.add_subj_proj_cb.clear()
        self.add_subj_proj_cb.addItem("", None)
        for pid, pname, comp_id in self.projects:
            if comp_id == cid:
                self.add_subj_proj_cb.addItem(pname, pid)

    
    def add_subject(self):
        name = self.add_subj_le.text().strip()
        proj_id = self.add_subj_proj_cb.currentData()
        if not name or proj_id is None:
            QMessageBox.warning(self, "Hata", "Lütfen konu adı ve projeyi seçin.")
            return
        conn = self.db_connect(); cur = conn.cursor()
        try:
            cur.execute(
                "INSERT INTO subjects (name, project_id) VALUES (%s, %s)",
                (name, proj_id)
            )
            conn.commit()
            QMessageBox.information(self, "Tamam", "Konu eklendi.")
        except mysql.connector.IntegrityError:
            QMessageBox.warning(self, "Hata", "Bu konu zaten var.")
        finally:
            cur.close(); conn.close()
        self.add_subj_le.clear()
        self.load_lookups()
        self.refresh_lookup_widgets()


    def delete_subject(self):
        sid = self.del_subj_cb.currentData()
        if sid is None:
            return
        ans = QMessageBox.question(
            self, "Konu Silme",
            "Seçili konudaki tüm ilişkili evraklardan konu bilgisi kaldırılacak, silmek istiyor musunuz?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if ans != QMessageBox.StandardButton.Yes:
            return
        conn = self.db_connect(); cur = conn.cursor()
        cur.execute("UPDATE documents SET subject_id=NULL WHERE subject_id=%s", (sid,))
        cur.execute("DELETE FROM subjects WHERE id=%s", (sid,))
        conn.commit()
        cur.close(); conn.close()
        self.load_lookups()
        self.refresh_lookup_widgets()


    def choose_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Evrak Dosyası Seç")
        if path:
            self.file_path_le.setText(path)
    def choose_target_folder(self):
        start = self.cfg.get('archive_root', os.path.expanduser("~"))
        options = (
            QFileDialog.Option.ShowDirsOnly 
            | QFileDialog.Option.DontUseNativeDialog
        )
        path = QFileDialog.getExistingDirectory(
            self,
            "Kayıt Klasörü Seçin",
            start,
            options
        )
        if path:
            self.target_folder_le.setText(path)


    def on_save(self):
        src = self.file_path_le.text().strip()
        if not self.edit_mode and not os.path.isfile(src):
            QMessageBox.warning(self, "Hata", "Lütfen geçerli bir dosya seçin.")
            return

        folder = self.target_folder_le.text().strip()
        if not os.path.isdir(folder):
            QMessageBox.warning(self, "Hata", "Lütfen geçerli bir kayıt klasörü seçin.")
            return

        if self.company_cb.currentData() is None:
            QMessageBox.warning(self, "Eksik Alan", "Lütfen şirket seçin.")
            return
        if self.project_cb.currentData() is None:
            QMessageBox.warning(self, "Eksik Alan", "Lütfen proje seçin.")
            return
        if self.subject_cb.currentData() is None:
            QMessageBox.warning(self, "Eksik Alan", "Lütfen konu seçin.")
            return

        an = self.archive_name_le.text().strip()
        subj = self.subject_cb.currentText().strip()
        subject_id = self.subject_cb.currentData()
        dtc = self.doc_type_cb.currentData()
        dd = self.doc_date_de.date().toString("yyyy-MM-dd")
        resp = self.resp_cb.currentData()
        snd = self.sender_le.text().strip()
        rec = self.recipient_le.text().strip()
        if not all([an, subj, dtc, resp, snd, rec]):
            QMessageBox.warning(self, "Eksik Alan", "Zorunlu(*) tüm alanları doldurun.")
            return

        comp = self.company_cb.currentData()
        proj = self.project_cb.currentData()
        sd = self.start_date_de.date().toString("yyyy-MM-dd")
        ed = self.end_date_de.date().toString("yyyy-MM-dd")
        comm = self.comment_te.toPlainText().strip() or None

        conn = self.db_connect()
        cur = conn.cursor()

        if self.edit_mode:
            tn = self.edit_tracking_no
            cur.execute("SELECT file_path FROM documents WHERE tracking_number=%s", (tn,))
            old_path = cur.fetchone()[0]
            ext = os.path.splitext(old_path)[1]
            filename = f"{an}{ext}"
            dst = os.path.join(folder, filename)
            if src and os.path.isfile(src):
                try:
                    shutil.copy2(src, dst)
                except PermissionError:
                    QMessageBox.critical(self, "Dosya Kullanımda", "Seçilen dosya şu anda başka bir program tarafından kullanılıyor. Lütfen kapatın ve tekrar deneyin.")
                    return
            else:
                dst = old_path

            sql = """
                UPDATE documents SET
                    archive_name=%s, subject=%s, subject_id=%s, doc_type=%s, company=%s,
                    project=%s, doc_date=%s, start_date=%s, end_date=%s,
                    responsible=%s, sender=%s, recipient=%s,
                    comment=%s, file_path=%s
                WHERE tracking_number=%s
            """
            vals = (
                an, subj, subject_id, dtc, comp,
                proj, dd, sd, ed,
                resp, snd, rec,
                comm, dst,
                tn
            )
            cur.execute(sql, vals)
            conn.commit()
            QMessageBox.information(self, "Güncellendi", f"{tn} kodlu evrak güncellendi.")
        else:
            tn = self.generate_tracking_no()
            ext = os.path.splitext(src)[1]
            filename = f"{an}{ext}"
            dst = os.path.join(self.target_folder_le.text(), f"{an}{ext}")
            shutil.copy2(src, dst)

            sql = """
                INSERT INTO documents
                (tracking_number, doc_type, company, project,
                doc_date, start_date, end_date,
                responsible, sender, recipient,
                archive_name, subject, comment, related_id, file_path, subject_id)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            """
            vals = (
                tn, dtc, comp, proj,
                dd, sd, ed,
                resp, snd, rec,
                an, subj, comm, None, dst, subject_id
            )
            cur.execute(sql, vals)

            parent_id = self.link_parent_cb.currentData()
            if parent_id is not None:
                cur.execute(
                    "INSERT INTO document_links (parent_id, child_id, link_type) VALUES (%s, %s, %s)",
                    (parent_id, cur.lastrowid, 'reply')
                )

            conn.commit()
            QMessageBox.information(self, "Kaydedildi", f"Evrak eklendi: {tn}")

        cur.close()
        conn.close()

        self.clear_form()
        self.load_lookups()
        self.refresh_lookup_widgets()


    def search_records(self):
        # 1) OFFSET ve LIMIT hesapla
        offset = self.current_page * self.page_size
        limit  = self.page_size

        # 2) Koşullar ve parametreler
        conds, params = [], []

        # Arşiv Adlandırması filtresi
        a = self.search_archive_le.text().strip()
        if a:
            conds.append("d.archive_name LIKE %s")
            params.append(f"%{a}%")

        # Takip No filtresi
        t = self.search_tracking_le.text().strip()
        if t:
            conds.append("d.tracking_number LIKE %s")
            params.append(f"%{t}%")

        # Evrak Tipi filtresi
        dt = self.search_doc_type_cb.currentData()
        if dt:
            conds.append("d.doc_type = %s")
            params.append(dt)

        # Şirket filtresi
        c = self.search_company_cb.currentData()
        if c is not None:
            conds.append("d.company = %s")
            params.append(c)

        # Proje filtresi
        p = self.search_project_cb.currentData()
        if p is not None:
            conds.append("d.project = %s")
            params.append(p)

        # Sorumlu filtresi
        r = self.search_res_cb.currentData()
        if r is not None:
            conds.append("d.responsible = %s")
            params.append(r)

        # Tarih aralığı filtresi
        df = self.search_date_from.date().toString("yyyy-MM-dd")
        dt2= self.search_date_to.date().toString("yyyy-MM-dd")
        conds.append("d.doc_date BETWEEN %s AND %s")
        params.extend([df, dt2])

        # 3) Toplam kayıt sayısını al (sayfalama için)
        count_sql = "SELECT COUNT(*) FROM documents d"
        if conds:
            count_sql += " WHERE " + " AND ".join(conds)

        conn = self.db_connect()
        cur  = conn.cursor()
        cur.execute(count_sql, tuple(params))
        self.total_records = cur.fetchone()[0] or 0

        # 4) Asıl veri sorgusu: JOIN ile isimleri çekiyoruz
        data_sql = """
            SELECT
                d.tracking_number,
                d.archive_name,
                d.doc_type,
                co.name   AS company,
                pr.name   AS project,
                d.doc_date,
                u.name    AS responsible,
                d.sender,
                d.recipient,
                d.comment
            FROM documents d
            LEFT JOIN companies co ON d.company     = co.id
            LEFT JOIN projects  pr ON d.project     = pr.id
            LEFT JOIN users     u  ON d.responsible = u.id
        """
        if conds:
            data_sql += " WHERE " + " AND ".join(conds)

        data_sql += " ORDER BY d.created_at DESC"
        data_sql += " LIMIT %s OFFSET %s"
        params.extend([limit, offset])

        cur.execute(data_sql, tuple(params))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        # 5) Sonuçları tabloya doldur
        self.results_table.setRowCount(len(rows))
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                text = "" if val is None else str(val)
                self.results_table.setItem(i, j, QTableWidgetItem(text))

        # 6) Sayfa bilgisi güncelle ve butonları ayarla
        total_pages = max(1, math.ceil(self.total_records / self.page_size))
        self.page_lbl.setText(f"Sayfa: {self.current_page+1} / {total_pages}")
        self.prev_btn.setEnabled(self.current_page > 0)
        self.next_btn.setEnabled(self.current_page < total_pages-1)


    def add_responsible(self):
        name = self.add_res_le.text().strip()
        if not name:
            QMessageBox.warning(self, "Hata", "Lütfen sorumlu kişinin tam adını girin.")
            return
        username = compute_username(name)
        password = DEFAULT_PASSWORD
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO users (name, username, password) VALUES (%s, %s, %s)",
            (name, username, password)
        )
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(
            self, "Tamam",
            f"Sorumlu eklendi:\n  Ad Soyad: {name}\n  Kullanıcı Adı: {username}\n  Şifre: {password}"
        )
        self.add_res_le.clear()
        self.load_lookups()
        self.refresh_lookup_widgets()

    def delete_responsible(self):
        id_ = self.del_res_cb.currentData()
        if id_ is None:
            return
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("DELETE FROM users WHERE id=%s", (id_,))
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(self, "Tamam", "Sorumlu silindi.")
        self.load_lookups()
        self.refresh_lookup_widgets()

    def add_company(self):
        name = self.add_company_le.text().strip()
        if not name:
            QMessageBox.warning(self, "Hata", "Şirket adı girin.")
            return
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("INSERT INTO companies (name) VALUES (%s)", (name,))
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(self, "Tamam", "Şirket eklendi.")
        self.add_company_le.clear()
        self.load_lookups()
        self.refresh_lookup_widgets()

    def delete_company(self):
        id_ = self.del_company_cb.currentData()
        if id_ is None:
            return
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("DELETE FROM companies WHERE id=%s", (id_,))
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(self, "Tamam", "Şirket silindi.")
        self.load_lookups()
        self.refresh_lookup_widgets()

    def add_project(self):
        name    = self.add_proj_le.text().strip()
        comp_id = self.add_proj_company_cb.currentData()
        if not name or comp_id is None:
            QMessageBox.warning(self, "Hata", "Proje ve şirket seçin.")
            return
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("INSERT INTO projects (name, company_id) VALUES (%s,%s)", (name, comp_id))
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(self, "Tamam", "Proje eklendi.")
        self.add_proj_le.clear()
        self.load_lookups()
        self.refresh_lookup_widgets()

    def delete_project(self):
        id_ = self.del_project_cb.currentData()
        if id_ is None:
            return
        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("DELETE FROM projects WHERE id=%s", (id_,))
        conn.commit(); cur.close(); conn.close()
        QMessageBox.information(self, "Tamam", "Proje silindi.")
        self.load_lookups()
    def preview_document(self, row, col):
        # tabloya tıklanan dosyanın path’ini al
        tn = self.results_table.item(row, 0).text()
        conn = self.db_connect(); cur = conn.cursor()
        cur.execute("SELECT file_path FROM documents WHERE tracking_number=%s", (tn,))
        res = cur.fetchone()
        cur.close(); conn.close()
        self.refresh_lookup_widgets()

        if not res or not os.path.exists(res[0]):
            QMessageBox.warning(self, "Hata", "Dosya bulunamadı.")
            return
        path = res[0]
        ext = os.path.splitext(path)[1].lower()

        # Önce tüm preview widget’lerini gizle
        for w in (self.pdf_view, self.text_view, self.xlsx_table, self.image_view):
            w.hide()

        if ext == ".pdf":
            # PDF görüntüle
            self.pdf_doc.load(path)
            self.pdf_view.show()

        elif ext in (".docx",):
            # DOCX -> Text
            doc = docx.Document(path)
            text = "\n\n".join(p.text for p in doc.paragraphs)
            self.text_view.setPlainText(text)
            self.text_view.show()

        elif ext in (".xlsx", ".xls"):
            # Excel -> QTableWidget
            wb = openpyxl.load_workbook(path, data_only=True)
            ws = wb.active
            rows = list(ws.rows)
            self.xlsx_table.clear()
            self.xlsx_table.setRowCount(len(rows))
            self.xlsx_table.setColumnCount(len(rows[0]))
            for i, row_cells in enumerate(rows):
                for j, cell in enumerate(row_cells):
                    val = "" if cell.value is None else str(cell.value)
                    self.xlsx_table.setItem(i, j, QTableWidgetItem(val))
            self.xlsx_table.show()

        elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".gif"):
            # Görsel önizleme
            pix = QPixmap(path)
            scaled = pix.scaled(
                self.preview_stack.size(),
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
            self.image_view.setPixmap(scaled)
            self.image_view.show()

        else:
            QMessageBox.information(
                self, "Önizleme Yok",
                "Bu dosya tipi için önizleme desteklenmiyor."
            )


    def build_notifications_tab(self):
            layout = QVBoxLayout()

            # Tablo başlıkları: ID, Takip No, Arşiv Ad., Bitiş Tarihi, Kalan Gün, İşlemler
            self.notif_table = QTableWidget()
            headers = ["ID","Takip No","Arşiv Ad.","Bitiş Tarihi","Kalan Gün","İşlemler"]
            self.notif_table.setColumnCount(len(headers))
            self.notif_table.setHorizontalHeaderLabels(headers)

            # Yenile düğmesi
            btn = QPushButton("Yenile")
            btn.clicked.connect(self.load_notifications)

            layout.addWidget(btn)
            layout.addWidget(self.notif_table)

            w = QWidget()
            w.setLayout(layout)
            return w
    def load_notifications(self):
        conn = self.db_connect()
        cur  = conn.cursor()
        cur.execute("""
        SELECT id, tracking_number, archive_name, end_date,
               DATEDIFF(end_date, CURDATE()) AS days_left
          FROM documents
         WHERE end_date IS NOT NULL
           AND notify_suppressed = 0
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()

        self.notif_table.setRowCount(len(rows))
        for i, (doc_id, tn, an, ed, days_left) in enumerate(rows):
            self.notif_table.setItem(i, 0, QTableWidgetItem(str(doc_id)))
            self.notif_table.setItem(i, 1, QTableWidgetItem(tn))
            self.notif_table.setItem(i, 2, QTableWidgetItem(an or ""))
            self.notif_table.setItem(i, 3, QTableWidgetItem(str(ed)))
            self.notif_table.setItem(i, 4, QTableWidgetItem(str(days_left)))

            # Her satıra "Bir Daha Gösterme" butonu ekle
            btn = QPushButton("Sustur")
            btn.clicked.connect(partial(self.suppress_notification, doc_id))
            self.notif_table.setCellWidget(i, 5, btn)
    def suppress_notification(self, doc_id):
        conn = self.db_connect()
        cur  = conn.cursor()
        cur.execute("UPDATE documents SET notify_suppressed=1 WHERE id=%s", (doc_id,))
        conn.commit()
        cur.close()
        conn.close()
        QMessageBox.information(self, "Tamam", "Bu bildirim bir daha gösterilmeyecek.")
        self.load_notifications()

            
    def check_deadlines(self):
        conn = self.db_connect()
        cur  = conn.cursor()
        cur.execute("""
            SELECT id, tracking_number, archive_name, end_date
              FROM documents
             WHERE end_date IS NOT NULL
               AND notify_suppressed = 0
               AND DATEDIFF(end_date, CURDATE()) BETWEEN 0 AND 7
        """)
        upcoming = cur.fetchall()
        cur.close()
        conn.close()

        for doc_id, tn, an, ed in upcoming:
            days_left = (ed - datetime.date.today()).days
            QMessageBox.information(
                self, "Süre Uyarısı",
                f"Evrak {tn} ({an}) için bitişe {days_left} gün kaldı: {ed}"
            )

    def showEvent(self, event):
        super().showEvent(event)
        # only run once, the first time the window is shown
        QTimer.singleShot(0, self.check_deadlines)
    def on_next_page(self):
        self.current_page += 1
        self.search_records()

    def on_prev_page(self):
        self.current_page -= 1
        self.search_records()
    def open_context_menu(self, pos):
        idx = self.results_table.indexAt(pos)
        if not idx.isValid():
            return
        row = idx.row()

        menu = QMenu(self)
        act_preview  = menu.addAction("Önizleme")
        act_open_loc = menu.addAction("Dosya Konumunu Aç")
        act_edit     = menu.addAction("Düzenle")
        act_delete   = menu.addAction("Sil")

        action = menu.exec(self.results_table.viewport().mapToGlobal(pos))
        if action == act_preview:
            # 0. sütun tıklanmış gibi preview çalıştırıyoruz
            self.preview_document(row, 0)
        elif action == act_open_loc:
            self.open_file_location(row)
        elif action == act_edit:
            self.edit_record(row)
        elif action == act_delete:
            self.delete_record(row)
    def open_file_location(self, row):
        # 1) Takip numarasını al
        tn = self.results_table.item(row, 0).text()

        # 2) DB'den dosya yolunu çek
        conn = self.db_connect()
        cur  = conn.cursor()
        cur.execute("SELECT file_path FROM documents WHERE tracking_number=%s", (tn,))
        res = cur.fetchone()
        cur.close()
        conn.close()

        if not res or not res[0]:
            QMessageBox.warning(self, "Hata", "Dosya bulunamadı.")
            return

        # 3) Klasör yolunu normalize et
        folder = os.path.dirname(res[0])
        folder = os.path.normpath(folder)

        if not os.path.exists(folder):
            QMessageBox.warning(self, "Hata", f"Dizin bulunamadı:\n{folder}")
            return

        # 4) Windows Explorer'da aç
        import subprocess, sys
        if sys.platform.startswith("win"):
            subprocess.Popen(["explorer", folder])
        else:
            # Diğer platformlarda xdg-open (Linux) veya open (macOS)
            opener = "xdg-open" if sys.platform.startswith("linux") else "open"
            subprocess.Popen([opener, folder])


    def edit_record(self, row):
        tn = self.results_table.item(row, 0).text()

        conn = self.db_connect()
        cur = conn.cursor()
        cur.execute("""
            SELECT archive_name, doc_type, company, project, doc_date,
                responsible, sender, recipient, comment
            FROM documents WHERE tracking_number=%s
        """, (tn,))
        rec = cur.fetchone()
        cur.close()
        conn.close()
        if not rec:
            QMessageBox.warning(self, "Hata", "Kayıt bulunamadı.")
            return

        dlg = QDialog(self)
        dlg.setWindowTitle(f"{tn} Düzenle")
        form = QFormLayout(dlg)

        an_le = QLineEdit(rec[0])
        form.addRow("Arşiv Adı", an_le)

        dt_cb = QComboBox()
        dt_cb.addItem("Resmi Evrak", "official")
        dt_cb.addItem("Taşınmaz Evrak", "immovable")
        dt_cb.addItem("Proje", "project")
        dt_cb.setCurrentIndex(dt_cb.findData(rec[1]))
        form.addRow("Evrak Tipi", dt_cb)

        co_cb = QComboBox()
        for cid, cname in self.companies:
            co_cb.addItem(cname, cid)
        co_cb.setCurrentIndex(co_cb.findData(rec[2]))
        form.addRow("Şirket", co_cb)

        pr_cb = QComboBox()
        for pid, pname, _ in self.projects:
            pr_cb.addItem(pname, pid)
        pr_cb.setCurrentIndex(pr_cb.findData(rec[3]))
        form.addRow("Proje", pr_cb)

        date_de = QDateEdit(QDate.fromString(rec[4].strftime("%Y-%m-%d"), "yyyy-MM-dd"))
        date_de.setCalendarPopup(True)
        form.addRow("Tarih", date_de)

        resp_cb = QComboBox()
        for uid, uname in self.responsibles:
            resp_cb.addItem(uname, uid)
        resp_cb.setCurrentIndex(resp_cb.findData(rec[5]))
        form.addRow("Sorumlu", resp_cb)

        sender_le = QLineEdit(rec[6])
        form.addRow("Gönderen", sender_le)

        recipient_le = QLineEdit(rec[7])
        form.addRow("Alıcı", recipient_le)

        comment_le = QLineEdit(rec[8] or "")
        form.addRow("Yorum", comment_le)

        btns = QHBoxLayout()
        save_btn = QPushButton("Kaydet")
        cancel_btn = QPushButton("İptal")
        btns.addWidget(save_btn)
        btns.addWidget(cancel_btn)
        form.addRow(btns)

        def on_save_edit():
            new_vals = (
                an_le.text().strip(),
                dt_cb.currentData(),
                co_cb.currentData(),
                pr_cb.currentData(),
                date_de.date().toString("yyyy-MM-dd"),
                resp_cb.currentData(),
                sender_le.text().strip(),
                recipient_le.text().strip(),
                comment_le.text().strip() or None,
                tn
            )
            sql = """
                UPDATE documents SET
                    archive_name=%s, doc_type=%s, company=%s, project=%s,
                    doc_date=%s, responsible=%s, sender=%s, recipient=%s,
                    comment=%s WHERE tracking_number=%s
            """
            conn = self.db_connect()
            cur = conn.cursor()
            cur.execute(sql, new_vals)
            conn.commit()
            cur.close()
            conn.close()

            QMessageBox.information(self, "Güncellendi", f"{tn} güncellendi.")
            dlg.accept()
            self.search_records()

        save_btn.clicked.connect(on_save_edit)
        cancel_btn.clicked.connect(dlg.reject)

        dlg.exec()

    def delete_record(self, row):
        tn = self.results_table.item(row, 0).text()

        # Şifre sor
        pw, ok = QInputDialog.getText(
            self, "Şifre Doğrulama",
            "Silmek için şifrenizi girin:",
            QLineEdit.EchoMode.Password
        )
        if not ok or not pw:
            return

        # Şifreyi DB’den kontrol et (login yaparken saklanan username)
        u = self.current_username  # login’dan gelen kullanıcı adı
        conn = self.db_connect(); cur = conn.cursor()
        cur.execute("SELECT password FROM users WHERE username=%s", (u,))
        row_pw = cur.fetchone()
        cur.close(); conn.close()

        if not row_pw or row_pw[0] != pw:
            QMessageBox.critical(self, "Hata", "Şifre yanlış.")
            return

        # Onay kutusu
        ans = QMessageBox.question(
            self, "Silme Onayı",
            f"{tn} numaralı evrak silinsin mi?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if ans != QMessageBox.StandardButton.Yes:
            return

        # Silme işlemi
        conn = self.db_connect(); cur = conn.cursor()
        cur.execute("DELETE FROM documents WHERE tracking_number=%s", (tn,))
        conn.commit(); cur.close(); conn.close()

        # Tablo satırını kaldır
        self.results_table.removeRow(row)
        QMessageBox.information(self, "Tamam", "Evrak silindi.")
    def refresh_all(self):
        self.load_lookups()
        self.refresh_lookup_widgets()
        self.search_records()
        self.load_notifications()
        QMessageBox.information(self, "Yenilendi", "Tüm veriler ve görünüm güncellendi.")

    def seed_example_workflow(self):
        conn = self.db_connect()
        cur = conn.cursor()

        # Şerket ekle
        cur.execute("INSERT INTO companies (name) VALUES ('Test Şerketi')")
        company_id = cur.lastrowid

        # Proje ekle
        cur.execute("INSERT INTO projects (name, company_id) VALUES ('Test Projesi', %s)", (company_id,))
        project_id = cur.lastrowid

        # Konu ekle
        cur.execute("INSERT INTO subjects (name, project_id) VALUES ('Test Konusu', %s)", (project_id,))
        subject_id = cur.lastrowid

        # Belgeler
        tracking_numbers = []
        end_date = (datetime.date.today() + datetime.timedelta(days=5)).strftime('%Y-%m-%d')

        for i in range(1, 4):
            tn = f"TEST-{i:03}"
            cur.execute("""
                INSERT INTO documents (
                    tracking_number, archive_name, doc_type, company,
                    project, subject_id, doc_date, end_date,
                    responsible, sender, recipient
                ) VALUES (%s, %s, %s, %s, %s, %s, CURDATE(), %s, NULL, 'Gönderen', 'Alıcı')
            """, (
                tn, f"Arşiv {i}", 'official', company_id, project_id, subject_id, end_date
            ))
            tracking_numbers.append(cur.lastrowid)

        # Zincir bağlantı: 1 → 2 → 3
        cur.execute("INSERT INTO document_links (parent_id, child_id) VALUES (%s, %s)", (tracking_numbers[0], tracking_numbers[1]))
        cur.execute("INSERT INTO document_links (parent_id, child_id) VALUES (%s, %s)", (tracking_numbers[1], tracking_numbers[2]))

        conn.commit()
        cur.close()
        conn.close()

        QMessageBox.information(self, "Hazır", "Test evrakları ve akış zinciri oluşturuldu.")


    def delete_test_data(self):
        conn = self.db_connect()
        cur = conn.cursor()
        try:
            cur.execute("DELETE FROM document_links WHERE child_id IN (SELECT id FROM documents WHERE tracking_number LIKE 'TEST-%')")
            cur.execute("DELETE FROM documents WHERE tracking_number LIKE 'TEST-%'")
            cur.execute("DELETE FROM subjects WHERE name = 'Test Konusu'")
            cur.execute("DELETE FROM projects WHERE name = 'Test Projesi'")
            cur.execute("DELETE FROM companies WHERE name = 'Test Şirketi'")
            conn.commit()
            QMessageBox.information(self, "Silindi", "Test verileri silindi.")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Silme sırasında hata: {e}")
        finally:
            cur.close()
            conn.close()


if __name__ == "__main__":

    app = QApplication(sys.argv)

    # Tema ayarları
    app.setStyle("Fusion")
    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor("#f5f5f5"))
    palette.setColor(QPalette.ColorRole.Button, QColor("#ffffff"))
    palette.setColor(QPalette.ColorRole.ButtonText, QColor("#333333"))
    palette.setColor(QPalette.ColorRole.WindowText, QColor("#222222"))
    palette.setColor(QPalette.ColorRole.Text, QColor("#222222"))
    palette.setColor(QPalette.ColorRole.Highlight, QColor("#0078d7"))
    palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))
    app.setPalette(palette)
    font = QFont("Segoe UI", 10)
    app.setFont(font)

    app.setStyleSheet("""
        QTabWidget::pane {
            border: 1px solid #c4c4c4;
            border-radius: 4px;
        }
        QTabBar::tab {
            background: #e0e0e0;
            border: 1px solid #c4c4c4;
            padding: 8px 16px;
            border-top-left-radius: 4px;
            border-top-right-radius: 4px;
            min-width: 80px;
        }
        QTabBar::tab:selected {
            background: white;
            border-bottom-color: white;
        }
        QPushButton {
            background-color: #0078d7;
            color: white;
            border: none;
            padding: 6px 12px;
            border-radius: 4px;
        }
        QPushButton:hover {
            background-color: #005a9e;
        }
        QLineEdit, QPlainTextEdit, QDateEdit, QComboBox {
            border: 1px solid #c4c4c4;
            border-radius: 4px;
            padding: 4px;
            background-color: white;
        }
        QTableWidget {
            gridline-color: #ddd;
            selection-background-color: #0078d7;
            selection-color: #fff;
        }
        QComboBox QAbstractItemView {
            background-color: white;
            selection-background-color: #0078d7;
            selection-color: white;
        }
        QComboBox QAbstractItemView::item:hover {
            background-color: #005a9e;
            color: white;
        }
    """)

    # Logo ve config
    app_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    logo_path = os.path.join(app_dir, 'resources', 'logo.ico')
    app.setWindowIcon(QIcon(logo_path))

    with open(os.path.join(app_dir, 'config.json'), encoding='utf8') as f:
        cfg = json.load(f)

    cfg['files_dir'] = os.path.join(app_dir, cfg['files_dir'])
    mc = cfg['mysql']
    conn_params = {
        'host': mc['host'], 'port': mc['port'],
        'user': mc['user'], 'password': mc['password'],
        'database': mc['database'], 'charset': 'utf8mb4'
    }

    login_dialog = SecureLoginDialog(conn_params)
    if login_dialog.exec() == QDialog.DialogCode.Accepted:
        user = login_dialog.current_user
        w = MainWindow(cfg, conn_params, user)
        w.setWindowTitle("HAN Holding - Evrak Takibi Sistemi")
        w.setWindowIcon(QIcon(logo_path))
        w.resize(1000, 700)
        w.show()
        sys.exit(app.exec())
    else:
        sys.exit(0)

